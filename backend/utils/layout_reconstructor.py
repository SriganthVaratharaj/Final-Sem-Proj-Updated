"""
backend/utils/layout_reconstructor.py

Digital Twin / Spatial Reconstruction Utility
=============================================
Takes OCR results (text + bounding boxes) and reconstructs the document's 
layout into a plain text grid. This allows for a "digital copy" that
preserves the original positioning of headers, tables, and footers.

Advantage: Converts unclear/artistic fonts into clear digital text 
while maintaining the exact bill format.
"""
from __future__ import annotations
import logging
from typing import Any

logger = logging.getLogger(__name__)

def reconstruct_spatial_text(
    texts: list[str], 
    boxes: list[list[list[float]]], 
    image_width: int = 1000, 
    image_height: int = 1000,
    grid_width: int = 100,
    grid_height: int = 60
) -> str:
    """
    Creates a spatial text map (digital twin) of the invoice.
    
    Args:
        texts: List of recognized strings
        boxes: List of 4-point bounding boxes [[[x1,y1], [x2,y2], [x3,y3], [x4,y4]], ...]
        image_width/height: Original image dimensions for normalization
        grid_width/height: Resolution of the text grid (columns/rows)
        
    Returns:
        A multiline string representing the invoice layout.
    """
    if not texts or not boxes:
        return ""

    # Create empty grid (list of lists of spaces)
    grid = [[" " for _ in range(grid_width)] for _ in range(grid_height)]

    for text, box in zip(texts, boxes):
        if not text or not box or len(box) < 4:
            continue
            
        # 1. Calculate center and dimensions in original pixels
        if isinstance(box[0], (list, tuple)):
            # Nested list: [[x1,y1], [x2,y2], ...]
            x_coords = [p[0] for p in box]
            y_coords = [p[1] for p in box]
            xmin, xmax = min(x_coords), max(x_coords)
            ymin, ymax = min(y_coords), max(y_coords)
        else:
            # Flat list: [x1, y1, x2, y2]
            xmin, ymin, xmax, ymax = box[:4]
        
        # 2. Normalize to grid coordinates
        # Map [0, image_width] -> [0, grid_width-1]
        # Map [0, image_height] -> [0, grid_height-1]
        
        col_start = int((xmin / image_width) * grid_width)
        row = int(((ymin + ymax) / 2 / image_height) * grid_height)
        
        # Bound checks
        row = max(0, min(row, grid_height - 1))
        col_start = max(0, min(col_start, grid_width - 1))
        
        # 3. Place text into grid
        # If the slot is already taken, we append or find next available space
        # (Very basic collision handling)
        for i, char in enumerate(text):
            curr_col = col_start + i
            if curr_col < grid_width:
                # If we're overwriting something that isn't a space, 
                # maybe we keep the longer string or just overwrite
                grid[row][curr_col] = char

    # 4. Convert grid to string
    lines = []
    for row_data in grid:
        line = "".join(row_data).rstrip()
        # Only add non-empty lines
        if line:
            lines.append(line)
        else:
            # Keep empty lines to maintain vertical scale, but limit sequences
            if lines and lines[-1] != "":
                lines.append("")
                
    return "\n".join(lines)

def save_digital_twin(
    content: str, 
    output_dir: Any, 
    base_filename: str, 
    extension: str = "txt"
) -> str:
    """Saves the reconstructed layout to a file."""
    import os
    from datetime import datetime
    from pathlib import Path
    
    out_path = Path(output_dir)
    out_path.mkdir(parents=True, exist_ok=True)
    
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"{base_filename}_digital_twin_{timestamp}.{extension}"
    full_path = out_path / filename
    
    if extension == "txt":
        full_path.write_text(content, encoding="utf-8")
    elif extension == "docx":
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            # Use a monospaced font like Courier New to preserve the grid
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Courier New'
            font.size = Pt(8)
            
            # Set narrow margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Pt(36)
                section.bottom_margin = Pt(36)
                section.left_margin = Pt(36)
                section.right_margin = Pt(36)

            for line in content.splitlines():
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.add_run(line)
                
            doc.save(str(full_path))
        except ImportError:
            logger.warning("python-docx not installed. Falling back to .txt only.")
            return str(save_digital_twin(content, output_dir, base_filename, "txt"))
            
    return str(full_path)
